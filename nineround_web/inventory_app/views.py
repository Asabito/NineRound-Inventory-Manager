from django.shortcuts import render, redirect
from django.http import HttpResponse
from .models import Event, Inventory # import models
from django.db.models import Q, F
from .forms import EventForm, InventoryForm
# user auth
from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login, logout
from django.contrib import messages
from django.contrib.auth.decorators import login_required, user_passes_test

# barcode
import barcode
from barcode.writer import ImageWriter
import PIL
from PIL import Image
from django.core.files.temp import NamedTemporaryFile
from django.core import files
from io import BytesIO

# docx
from docx import Document
from docx.shared import Mm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
import os, shutil
from docx.shared import Pt
PATH = os.path.dirname(os.path.abspath(__file__))
import datetime as dt

# download
import magic
from pathlib import Path
BASE_DIR = Path(__file__).resolve().parent.parent
IMG_DIR = os.path.join(BASE_DIR, 'inventory_app', 'temp', 'img')
DOC_DIR = os.path.join(BASE_DIR, 'inventory_app', 'temp', 'doc')

# upload
import pandas as pd

# Create your views here.
def loginPage(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        try:
            User.objects.get(username=username)
        except:
            messages.error(request, 'Username or password is incorrect')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('events')
        else:
            messages.error(request, 'Username or password is incorrect')
    return render(request, 'inventory_app/login.html')


def logoutUser(request):
    logout(request)
    return redirect('loginPage')


guest_only = lambda u: u.role == 'Guest'
admin_only = lambda u: u.role == 'Admin'
superadmin_only = lambda u: u.role == 'Superadmin'
admin_superadmin_only = lambda u: u.role in ['Admin','Superadmin']


@login_required(login_url='loginPage')
def events(request):
    q = request.GET.get('q') if request.GET.get('q') != None else ''
    events = Event.objects.all().order_by('timestamp')
    events = Event.objects.filter(
        Q(nama__icontains=q) |
        Q(lokasi__icontains=q) |
        Q(status__icontains=q)
    ).order_by('timestamp')
    context = {'events': events}
    if request.method == 'POST' and request.POST.get("delete-event-button"):
        events_id_to_delete = request.POST.getlist("events_to_delete")
        events_to_delete = Event.objects.filter(id__in=events_id_to_delete)
        print('Events to be deleted: ',events_to_delete)
        events_to_delete.delete()
        # return None
    elif request.method == 'POST' and request.POST.get("new-event-button"):
        return redirect('new-event')
    
    date_today = dt.date.today()
    for event in events:
        if not (event.tanggal_mulai <= date_today <= event.tanggal_berakhir):
            Event.objects.filter(id=event.id).update(status='Selesai')
        # print(event.id, '<---')
        # print(event.tanggal_mulai, '<---')
        # print(event.tanggal_berakhir)
        # print(dt.date.today())

    return render(request, 'inventory_app/event.html', context=context)
    # return HttpResponse('Home Page')

@login_required(login_url='loginPage')
@user_passes_test(admin_superadmin_only)
def newEvent(request):
    form = EventForm()
    # buat list barcode apabila belum ada
    if 'barcode' not in request.session:
        request.session['barcode'] = []
    # apabila klik submit-button atau delete-item, maka
    

    if request.method == 'POST':
        form = EventForm(request.POST)
        
        if form.is_valid() and request.POST.get("additem-button"):
            if request.POST['barcode-input'] in Inventory.objects.values_list('id', flat=True) and request.POST['barcode-input'] not in request.session['barcode']:
                request.session['barcode'] += [request.POST['barcode-input']]
        elif form.is_valid() and request.POST.get("submit-button"):
            new_event = Event(nama=request.POST['nama'],
                              lokasi= request.POST['lokasi'],
                              tanggal_mulai = request.POST['tanggal_mulai'],
                              tanggal_berakhir = request.POST['tanggal_berakhir']
                              )
            new_event.save()

            # update the item's FK to the Event (one item can only exist in an event or warehouse)
            Inventory.objects.filter(id__in=request.session['barcode']).update(items_event_location=Event.objects.all().order_by('-timestamp').first().id)
            # request.session.flush()
            del request.session['barcode']
            return redirect('events')
        
        elif request.POST.get('cancel-button'):
            if request.session['barcode']:
                # request.session.flush()
                del request.session['barcode']
            return redirect('events')
        elif form.is_valid() and request.POST.get("delete-button"):
            deleted_id = request.POST.getlist('items_to_delete')
            for i in deleted_id:
                request.session['barcode'].remove(i)
    request.session.modified = True
    items = Inventory.objects.filter(id__in=request.session['barcode']).order_by('items', 'id')
    context = {'form':form, 'items':items}
    return render(request, 'inventory_app/new-event.html', context=context)

@login_required(login_url='loginPage')
@user_passes_test(admin_superadmin_only)
def eventDetail(request, pk):
    event_details = Inventory.objects.filter(items_event_location=pk).order_by('items', 'id') # query all items in Inventory, dimana events id sama dengan pk yang dipassing (many to many relationship). Python memberikan kemudahan bagi developer untuk melakukan lookup dengan menggunakan *namaTabelLain*__*kolomTabelLainTersebut*
    event = Event.objects.filter(id=pk)
    context = {'event_details':event_details, 'event':event}

    return render(request, 'inventory_app/event-detail.html', context=context)

@login_required(login_url='loginPage')
@user_passes_test(admin_superadmin_only)
def addItemsToEvent(request, pk):
    # additional note: addItemsToEvent dan deleteItemsFromEvent bisa dibuatkan decoratornya agar lebih hemat penulisan
    event = Event.objects.filter(id=pk)
    all_inventory_id = Inventory.objects.values_list('id', flat=True)
    
    if 'list_item' not in request.session:
        request.session['list_item'] = []
    
    id_items_in_event = Inventory.objects.filter(items_event_location=pk).values_list('id', flat=True)
    print('id_items_in_event: ',id_items_in_event)
    if request.method == 'POST':
        item_id = request.POST.get('tambah-item-textfield')
        if request.POST.get('additem-to-event-button') and (item_id not in id_items_in_event) and (item_id in all_inventory_id):
            # Jika id yang dimasukkan belum ada di event dan ada di inventory, maka ditambahkan ke session
            request.session['list_item'].append(item_id)
        elif request.POST.get('save-button'):
            Inventory.objects.filter(id__in=request.session['list_item']).update(items_event_location=pk)
            # request.session.flush()
            del request.session['list_item']
            return redirect('eventDetail', pk=pk)
        elif request.POST.get('cancel-button'):
            # request.session.flush()
            del request.session['list_item']
            return redirect('eventDetail', pk=pk)

    event_items = Inventory.objects.filter(id__in=request.session['list_item'])
    # session harus disave agar hasil modified tersimpan
    request.session.modified = True
    context = {
        'events':event,
        'event_items':event_items
    }
    return render(request, 'inventory_app/add-item-to-event.html', context=context)
    None

@login_required(login_url='loginPage')
@user_passes_test(superadmin_only)
def deleteItemsFromEvent(request, pk):
    event = Event.objects.filter(id=pk)
    all_inventory_id = Inventory.objects.values_list('id', flat=True)
    
    if 'list_item' not in request.session:
        request.session['list_item'] = []
    
    id_items_in_event = Inventory.objects.filter(items_event_location=pk).values_list('id', flat=True)
    if request.method == 'POST':
        item_id = request.POST.get('tambah-item-textfield')
        if request.POST.get('additem-to-event-button') and (item_id in id_items_in_event) and (item_id in all_inventory_id):
            # Jika id yang dimasukkan belum ada di event dan ada di inventory, maka ditambahkan ke session
            request.session['list_item'].append(item_id)
        elif request.POST.get('delete-button'):
            Inventory.objects.filter(id__in=request.session['list_item']).update(items_event_location=None)
            # request.session.flush()
            del request.session['list_item']

            return redirect('eventDetail', pk=pk)
        elif request.POST.get('cancel-button'):
            # request.session.flush()
            del request.session['list_item']
            return redirect('eventDetail', pk=pk)

    event_items = Inventory.objects.filter(id__in=request.session['list_item'])
    # session harus disave agar hasil modified tersimpan
    request.session.modified = True
    context = {
        'events':event,
        'event_items':event_items
    }
    return render(request, 'inventory_app/delete-item-from-event.html', context=context)

@login_required(login_url='loginPage')
@user_passes_test(admin_superadmin_only)
def stockChecking(request, pk):
    if request.method == 'POST':
        update_to_tersedia = request.POST.get('tersediaText') # refer to form dengan method POST, dan ambil entity yang memiliki ID 'tersediaText'
        update_to_terjual = request.POST.get('terjualText') # refer to form dengan method POST, dan ambil entity yang memiliki ID 'terjualText'
        reset_clicked = request.POST.get('resetButton')

        # kalau yang di check-in adalah Tersedia, maka update di EventItems dan Inventory
        if update_to_tersedia:
            Inventory.objects.filter(id=update_to_tersedia).update(item_last_status='Tersedia')

        # kalau yang di check-in adalah Terjual, maka update di EventItems dan Inventory
        elif update_to_terjual:
            Inventory.objects.filter(id=update_to_terjual).update(item_last_status='Terjual')

        if reset_clicked:
            print('yesss, it clicked <----')
            a=Inventory.objects.filter(items_event_location=pk, item_last_status='Tersedia').update(item_last_status='Tidak ada')
            print(a)
            print('done, it clicked <----')

    event_details = Inventory.objects.filter(items_event_location=pk).annotate(items_status = F('item_last_status')).order_by('items', 'id')
    total_count = Inventory.objects.filter(items_event_location=pk).count()
    terjual_count = Inventory.objects.filter(item_last_status='Terjual', items_event_location=pk).count()
    barang_tersedia_count = Inventory.objects.filter(item_last_status='Tersedia', items_event_location=pk).count()
    barang_tidak_ada_count = Inventory.objects.filter(item_last_status='Tidak ada', items_event_location=pk).count()

    events = Event.objects.filter(id=pk)
    context = {'event_details':event_details, 
               'events':events, 
               'total_count': total_count,
               'terjual_count':terjual_count,
               'barang_tersedia_count':barang_tersedia_count,
               'barang_tidak_ada_count':barang_tidak_ada_count,
               }
    return render(request, 'inventory_app/stock-checking.html', context=context)

@login_required(login_url='loginPage')
@user_passes_test(admin_superadmin_only)
def inventoryPage(request):
    items = Inventory.objects.all().order_by('items_group', 'id')
    if request.method == 'POST' and request.POST.get('add-button'):
        return redirect('inventoryAddItem')
    elif request.method == 'POST' and request.POST.get('delete-button'):
        return redirect('inventoryDeleteItem')
    elif request.method == 'POST' and request.POST.get('submit-upload'):
        # try:
            if 2621440 < request.FILES['upload-button'].size:
                print('error bro')
                raise MemoryError
            excelParser(request.FILES['upload-button'])
            del request.FILES['upload-button']
        # except:
        #     messages.error(request, 'File size is too much or id is used or wrong excel format!')

    context = {'items':items}

    return render(request, 'inventory_app/inventory.html', context=context)


def excelParser(file):
    df = pd.read_excel(file)
    data_to_be_inputted = []
    for d in df.values:
        print(d)
        new_item = Inventory(
                                id = str(d[0]).strip().upper(),
                                nama = d[1].strip().capitalize(),
                                keterangan = d[2].strip().capitalize(),
                                ukuran = str(d[3]).strip().upper(),
                                harga = str(d[4]).strip(),
                                items_group = d[5].strip().capitalize()
                            )
        new_item.full_clean()
        data_to_be_inputted.append(new_item)
    for vd in data_to_be_inputted:
        vd.save()


@login_required(login_url='loginPage')
@user_passes_test(admin_superadmin_only)
def inventoryAddItem(request):
    # raw form
    form = InventoryForm()

    # raise error apabila id duplicate
    raise_warning = False

    # add to session
    if "temp_inven" not in request.session:
        request.session["temp_inven"] = [] 

    # add button  
    if request.method =='POST' and request.POST.get("add-item-button") :
        current_id = request.POST["id"]

        # cek apakah ID yang dimasukkan sudah ada di SESSION atau belum
        for d in request.session["temp_inven"]:
            print(d['item_id'] == current_id)
            print('current id =', current_id)
            print('session id =',d['item_id'])
            if d['item_id'] == current_id:
                raise_warning = True
                break

        # cek apakah ID yang dimasukkan sudah ada di DATABASE atau belum
        if current_id in Inventory.objects.values_list('id', flat=True) :
            raise_warning = True
        
        # validasi form dengan input
        form = InventoryForm(request.POST)
        # apabila ID belum digunakan, dan form valid maka tambahkan input ke session
        if form.is_valid() and raise_warning == False:
            request.session["temp_inven"].append({
                'item_id' : current_id,
                'item_nama' : request.POST["nama"],
                'item_keterangan' : request.POST["keterangan"],
                'item_ukuran' : request.POST["ukuran"],
                'item_harga' : request.POST["harga"],
                'items_group' : request.POST["items_group"],
            })
            # save session
            request.session.modified = True

            return redirect('inventoryAddItem')

    # delete button
    if request.method == 'POST' and request.POST.get('delete-item-button'):
        to_be_deleted_items = request.POST.getlist('items_to_delete')
        print(to_be_deleted_items)
        ids_in_temp_inven = [x["item_id"] for x in request.session['temp_inven']]
        print(ids_in_temp_inven)
        del_index = []
        for item in to_be_deleted_items:
            del_index.append(ids_in_temp_inven.index(item))

        print('del_index: ',del_index)
        del_index.sort(reverse=True)
        for d_i in del_index:
            del request.session["temp_inven"][d_i]
        print('del_index: ',del_index)
        
        request.session.save()
                
        # for item in request.session['temp_inven']:
        #     if item['item_id'] in deleted_items:


    # submit button
    elif request.method == 'POST' and request.POST.get('submit-button'):
        # request.session.flush()
        for x in request.session["temp_inven"]:
            new_item = Inventory(
                                 id = x['item_id'],
                                 nama = x['item_nama'],
                                 keterangan = x['item_keterangan'],
                                 ukuran = x['item_ukuran'],
                                 harga = x['item_harga'],
                                 items_group = x['items_group'],
                                )
            new_item.save()
            # request.session.flush()
            if 'temp_inven' in request.session:
                del request.session["temp_inven"]
        return redirect('inventoryPage')
    
    # cancel button
    elif request.method == 'POST' and request.POST.get('cancel-button'):
        # request.session.flush()
        del request.session["temp_inven"]
        return redirect('inventoryPage')
    # items = Inventory.objects.all()

    context = {'form':form, 'items':request.session["temp_inven"], 'raise_warning':raise_warning}
    return render(request, 'inventory_app/inventory-add-item.html', context=context)


@login_required(login_url='loginPage')
@user_passes_test(superadmin_only)
def inventoryDeleteItem(request):
    # buat session apabila belum ada
    if 'temp_item_list_to_be_deleted_from_inventory' not in request.session:
        request.session['temp_item_list_to_be_deleted_from_inventory'] = []

    # apabila add item ditekan
    if request.method == 'POST' and request.POST.get('add-item-button'):
        current_id = request.POST.get('search-id')
        if current_id not in request.session['temp_item_list_to_be_deleted_from_inventory']:
            request.session['temp_item_list_to_be_deleted_from_inventory'].append(current_id)

    # show item yang dimasukkan melalui add item
    items_to_be_deleted = Inventory.objects.filter(id__in=request.session['temp_item_list_to_be_deleted_from_inventory']).order_by('items_group', 'id')
    
    if request.method == 'POST' and request.POST.get('submit-button'):
        items_to_be_deleted.delete()
        # request.session.flush()
        del request.session['temp_item_list_to_be_deleted_from_inventory']
        return redirect('inventoryPage')
    elif request.method == 'POST' and request.POST.get('cancel-button'):
        # request.session.flush()
        del request.session['temp_item_list_to_be_deleted_from_inventory']
        return redirect('inventoryPage')

    # save session modification
    request.session.modified = True


    context = {'items_to_be_deleted': items_to_be_deleted}
    return render(request, 'inventory_app/inventory-delete-item.html', context=context)

@login_required(login_url='loginPage')
@user_passes_test(admin_superadmin_only)
def barcodeGenerator(request):
    item = Inventory.objects.all().order_by('items_group', 'id')
    context = {'items':item}
    if request.POST.getlist('selected_items'):
        # generate barcode images
        barcodeImageGenerator(request.POST.getlist('selected_items'))
        # read and put barcode images to the docx
        barcodeDocxGenerator()
        
        # delete all barcode png in /temp/img/ folder
        folder = IMG_DIR
        for filename in os.listdir(folder):
            file_path = os.path.join(folder, filename)
            try:
                if (os.path.isfile(file_path) or os.path.islink(file_path)) :
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                print('Failed to delete %s. Reason: %s' % (file_path, e))
        
    return render(request, 'inventory_app/barcode.html', context=context)


def barcodeImageGenerator(list_of_id):
    """
    This function generates barcode image for each selected IDs
    """
    bar_class = barcode.get_barcode_class('code128')
    writer=ImageWriter()

    for identifier in list_of_id:
        code128 = bar_class(identifier, writer)
        print('--------------->', os.path.join(IMG_DIR, identifier))
        code128.save(os.path.join(IMG_DIR, identifier), {
            "module_width":.3, 
            "module_height":10, 
            "font_size": 8, 
            "text_distance": 3, 
            "quiet_zone": 1
            }) # save the originally generated image


def barcodeDocxGenerator():
    """
    This function put barcode images in temp folder into the docx
    """

    # initiate document
    document = Document()
    # page layout settings
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_height = Mm(20)
    section.page_width = Mm(40)
    section.left_margin = Mm(3)
    section.right_margin = Mm(3)
    section.top_margin = Mm(3)
    section.bottom_margin = Mm(0)

    # read barcode images
    for subdir, dirs, files in os.walk(IMG_DIR):
        for idx, file in enumerate(files):
            item_detail = Inventory.objects.filter(id=file[:-4])
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()

            # add barcode image to the docx
            r.add_picture(os.path.join(IMG_DIR, file), width=Mm(35))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            p.add_run(
                        f'{item_detail[0].id}\n{item_detail[0].nama}\n{item_detail[0].keterangan}\n{item_detail[0].ukuran}\nRp {int(item_detail[0].harga):,}'

                      ).font.size = Pt(6)
            # prevent adding new blank page 
            if idx != len(files)-1:
                r.add_break(WD_BREAK.PAGE)

    # delete existing docx
    if os.path.exists(os.path.join(DOC_DIR, 'generated_barcode.docx')):
        os.unlink(os.path.join(DOC_DIR, 'generated_barcode.docx'))
    # save the new docx to /temp/doc/
    document.save(os.path.join(DOC_DIR, 'generated_barcode.docx'))
    

def downloadFile(response):
    """
    This function provide file to be downloaded by user
    """
    docx_file = open(os.path.join(DOC_DIR, 'generated_barcode.docx'), "rb").read()
    content_type = magic.from_buffer(docx_file, mime=True)
    response = HttpResponse(docx_file, content_type=content_type);

    # pass file to the response(?)
    response['Content-Disposition'] = 'attachment; filename="generated_barcode.docx"'
    
    return response


# TODO:
# - DONE -  kasih validasi setiap kali delete 
# - DONE - barcode
# - add autofocus pada input fields untuk mempercepat penggunaan scanner

# untuk di bawah ini: database diubah dari many to many menjadi one to many (one event many items). 
# - ubah status tidak diketahui manjadi ... (brainstorming lg)
# - ubah database menjadi: ketika user memasukkan item ke sebuah event, di event lainnya akan terhapus. Dan status akan berubah menjadi dalam event
